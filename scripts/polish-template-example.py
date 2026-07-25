from __future__ import annotations

import argparse
import copy
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_cell_text(
    cell,
    text: str,
    *,
    centered: bool,
    bold: bool = False,
    font_size: float = 10.5,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_text(
    paragraph,
    text: str,
    *,
    font_name: str = "宋体",
    font_size: float = 10.5,
    bold: bool = False,
    first_line_indent: float | None = 21,
) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = (
        None if first_line_indent is None else Pt(first_line_indent)
    )
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman" if font_name != "Consolas" else "Consolas"
    run.font.size = Pt(font_size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def set_paragraph_shading(paragraph, fill: str | None) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    for shading in list(paragraph_properties.findall(qn("w:shd"))):
        paragraph_properties.remove(shading)
    if fill is not None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), fill)
        paragraph_properties.append(shading)


def fill_data_table(document: Document, rows: list[list[str]]) -> None:
    for table in document.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == ["序号", "观测项", "原始值", "处理值", "备注"]:
            if len(table.rows) < len(rows) + 1:
                raise ValueError("The data-analysis table does not have enough rows.")
            for row_index, values in enumerate(rows, start=1):
                for column_index, value in enumerate(values):
                    set_cell_text(
                        table.cell(row_index, column_index),
                        value,
                        centered=column_index in {0, 2, 3},
                    )
            return
    raise ValueError("Could not find the data-analysis table.")


def fill_review_score(document: Document, score: str) -> None:
    for table in document.tables:
        if not table.rows:
            continue
        score_cell = table.rows[0].cells[-1]
        if score_cell.text.strip() != "成绩":
            continue
        if any(row.cells[-1]._tc is not score_cell._tc for row in table.rows):
            continue

        score_cell.text = ""
        score_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_paragraph = score_cell.paragraphs[0]
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_paragraph.paragraph_format.space_before = 0
        label_paragraph.paragraph_format.space_after = Pt(6)
        label_run = label_paragraph.add_run("成绩")
        label_run.bold = True
        label_run.font.name = "Times New Roman"
        label_run.font.size = Pt(12)
        label_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")

        score_paragraph = score_cell.add_paragraph()
        score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_paragraph.paragraph_format.space_before = 0
        score_paragraph.paragraph_format.space_after = 0
        score_run = score_paragraph.add_run(score)
        score_run.bold = True
        score_run.font.name = "Times New Roman"
        score_run.font.size = Pt(18)
        score_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
        return
    raise ValueError("Could not find the merged review-score cell.")


def polish_code_notebook(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    code_start = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip().startswith("def merge_sort(values):")
    )
    item_four = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip().startswith("4. 测试空列表")
    )
    start_index = paragraphs.index(code_start)
    item_four_index = paragraphs.index(item_four)
    set_paragraph_text(
        code_start,
        "\n".join(
            [
                "def merge_sort(values):",
                "    if len(values) <= 1:",
                "        return values.copy()",
                "    middle = len(values) // 2",
                "    left = merge_sort(values[:middle])",
                "    right = merge_sort(values[middle:])",
                "    return merge(left, right)",
            ]
        ),
        font_name="Consolas",
        font_size=9,
        first_line_indent=None,
    )
    if "Code Block" in document.styles:
        code_start.style = document.styles["Code Block"]
    set_paragraph_shading(code_start, "F3F5F7")
    code_start.paragraph_format.line_spacing = 1.15
    code_start.paragraph_format.space_before = Pt(4)
    code_start.paragraph_format.space_after = Pt(6)
    for paragraph in paragraphs[start_index + 1 : item_four_index]:
        remove_paragraph(paragraph)

    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if (
            paragraph.style.name == "Code Block"
            and paragraph._element is not code_start._element
        ):
            paragraph.style = document.styles["Normal"]
            set_paragraph_shading(paragraph, None)
            set_paragraph_text(
                paragraph,
                text,
                first_line_indent=None if text[:2].rstrip(".").isdigit() else 21,
            )
        if text.startswith("pytest 共执行") or text == "命令输出：" or text.startswith(
            "数据规模从"
        ):
            paragraph.style = document.styles["Normal"]
            set_paragraph_shading(paragraph, None)
            set_paragraph_text(paragraph, text)
        if text.startswith("1000 items:"):
            paragraph.text = ""
            if "Command Output" in document.styles:
                paragraph.style = document.styles["Command Output"]
            set_paragraph_shading(paragraph, "EEF4EF")
            set_paragraph_text(
                paragraph,
                "\n".join(
                    [
                        "1000 items:     1.34 ms",
                        "5000 items:     7.92 ms",
                        "10000 items:   17.10 ms",
                        "50000 items:  102.63 ms",
                        "100000 items: 221.48 ms",
                    ]
                ),
                font_name="Consolas",
                font_size=9,
                first_line_indent=None,
            )
            paragraph.paragraph_format.line_spacing = 1.15


def replace_table_rows(table, rows: list[list[str]]) -> None:
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < len(rows):
        source_row = table.rows[-1]._tr
        table._tbl.append(copy.deepcopy(source_row))
    for row_index, values in enumerate(rows):
        if len(values) != len(table.rows[row_index].cells):
            raise ValueError("Replacement table row has an unexpected column count.")
        for column_index, value in enumerate(values):
            set_cell_text(
                table.cell(row_index, column_index),
                value,
                centered=True,
                bold=row_index == 0,
                font_size=10 if row_index == 0 else 9.5,
            )


def repeat_table_header(table) -> None:
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    if row_properties.find(qn("w:tblHeader")) is None:
        row_properties.append(OxmlElement("w:tblHeader"))


def set_heading(paragraph, text: str, *, level: int = 2) -> None:
    paragraph.text = ""
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")


def polish_project_dossier(document: Document) -> None:
    paragraphs = list(document.paragraphs)

    body_replacements = {
        "1. 项目背景": (
            "本项目以个人学习与日常事务管理为应用场景，目标是在不依赖服务器的条件下，"
            "完成任务录入、状态管理、组合筛选、截止日期提醒和备份恢复。应用采用 Kotlin "
            "与 Jetpack Compose 开发，任务数据由 Room 持久化，界面通过 ViewModel 和 "
            "StateFlow 统一管理状态。"
        ),
        "支持按标题": (
            "应用需支持新增、编辑、完成和删除任务，并可按全部、未完成、今天、已过期和"
            "已完成筛选；标题与备注支持关键词检索。所有写操作必须经过输入校验，空标题"
            "不得保存，提醒时间不得晚于截止时间。500 条数据下筛选操作应在 100 ms 内"
            "完成，普通增删改操作应在 300 ms 内给出反馈。"
        ),
        "1. 用户与使用场景": (
            "系统面向单机用户，不设置登录流程。用户可在课后录入作业和截止日期，在首页"
            "查看今日事项，完成任务后切换状态，并通过课程分类和关键词查找历史任务。任务"
            "实体包含标题、备注、分类、优先级、截止时间、提醒时间和完成状态，应用必须在"
            "无网络环境下完成全部基础操作。"
        ),
        "分类以普通字符串": (
            "首页同时显示未完成、今日到期和已过期数量，筛选条件与搜索词可以组合使用。"
            "完成任务后取消尚未触发的提醒，恢复为未完成时重新判断是否需要调度。备份导入"
            "前检查版本与字段完整性，解析或校验失败时不覆盖现有数据库；恢复过程在单一"
            "事务中完成。"
        ),
        "1. 总体结构": (
            "项目分为 presentation、domain 和 data 三层。presentation 包含 Compose 页面、"
            "导航、ViewModel 与 UI State；domain 包含实体、筛选条件、校验规则和用例；data "
            "包含 Room Entity、DAO、Repository、备份序列化与提醒调度器。ViewModel 只调用"
            "用例，Repository 负责实体转换并向上层暴露数据库 Flow。"
        ),
        "列表查询不为": (
            "任务管理模块覆盖新增、编辑、完成和删除；查询模块组合状态、日期与关键词条件；"
            "首页使用统计卡片、筛选标签、搜索框和 LazyColumn 任务列表。TaskListViewModel "
            "把任务流、搜索词和筛选条件合并为单一 TaskListUiState，避免列表与统计区域"
            "短暂不一致。"
        ),
        "界面只收集": (
            "保存任务时先提交数据库，再以 task-reminder-{id} 为唯一名称创建或替换 "
            "OneTimeWorkRequest；完成或删除任务时按同一名称取消。备份导出先读取一致性"
            "快照并写入临时文件，恢复时先在内存校验，再在事务中替换数据，完成后重新"
            "扫描未来提醒。"
        ),
        "导出时先读取": (
            "“今天”按设备本地时区的半开区间计算，逾期定义为截止时间早于当前时间且任务"
            "未完成。搜索词统一转为小写并去除首尾空格，标题或备注任一匹配即可返回。任务"
            "排序依次比较完成状态、逾期状态、优先级、截止时间和创建时间，保证结果稳定。"
        ),
        "1. 基础功能": (
            "应用能够新增、编辑、完成和删除任务。空标题会显示校验提示且保存按钮不可用。"
            "新建“数据库实验报告”并把截止时间设为当天 22:00 后，记录立即出现在“今天”"
            "列表；勾选完成后统计卡片同步更新，重新启动应用后状态保持。500 条测试数据的"
            "关键词筛选中位耗时为 18 ms。"
        ),
        "3. 提醒与恢复": (
            "任务提醒在退出应用后按时触发，编辑提醒时间后旧工作被替换且只收到一次通知；"
            "完成任务后通知取消。导出 500 条记录、清空数据库再恢复后，任务数量、完成状态"
            "和提醒数量与导出前一致。空标题、非法时间关系、组合筛选、提醒替换、备份拒绝"
            "与完整恢复等 12 项核心用例全部通过。"
        ),
    }
    for paragraph in paragraphs:
        stripped = paragraph.text.strip()
        for prefix, replacement in body_replacements.items():
            if stripped.startswith(prefix):
                set_paragraph_text(paragraph, replacement)
                break

    heading_replacements = {
        "3.1 开发工具与运行环境": "2.1 开发工具与运行环境",
        "3.2 数据与调试配置": "2.2 数据与调试配置",
        "2.1 业务需求分析": "3.1 业务需求分析",
        "2.2 可行性分析": "3.2 约束与可行性分析",
        "4.5 功能模块设计": "4.5 功能模块划分",
        "4.6 数据库设计": "4.6 数据库设计",
        "4.7 数据库表结构": "4.7 数据结构",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in heading_replacements:
            set_heading(paragraph, heading_replacements[text], level=2)

    captions = {
        "表4-1 功能模块表": "表4-1 功能模块表",
        "表4-2 数据库表": "表4-2 数据表说明",
        "表4-3 User用户表": "表4-3 task 任务表",
        "表4-4 BusinessData业务数据表": "表4-4 备份 JSON 根对象",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in captions:
            set_paragraph_text(
                paragraph,
                captions[text],
                font_size=10.5,
                first_line_indent=None,
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    replace_table_rows(
        document.tables[1],
        [
            ["功能模块", "包含子功能模块", "功能"],
            ["任务管理", "新增、编辑、完成、删除", "完成任务全生命周期管理"],
            ["查询筛选", "状态、日期、关键词", "组合筛选并同步统计结果"],
            ["提醒服务", "调度、替换、取消", "维护唯一后台提醒任务"],
            ["备份恢复", "导出、校验、事务恢复", "保证备份数据完整可恢复"],
        ],
    )
    replace_table_rows(
        document.tables[2],
        [
            ["序号", "数据库表", "数据表存储的内容"],
            ["1", "task", "存储任务内容、状态、日期与提醒信息"],
        ],
    )
    replace_table_rows(
        document.tables[3],
        [
            ["序号", "字段名", "字段类型", "说明", "备注"],
            ["1", "id", "INTEGER", "任务编号", "主键自增"],
            ["2", "title", "TEXT", "任务标题", "非空"],
            ["3", "note", "TEXT", "任务备注", "可空"],
            ["4", "category", "TEXT", "任务分类", "可空"],
            ["5", "priority", "TEXT", "优先级", "枚举"],
            ["6", "due_at", "INTEGER", "截止时间戳", "可空"],
            ["7", "remind_at", "INTEGER", "提醒时间戳", "可空"],
            ["8", "completed", "INTEGER", "完成状态", "0/1"],
        ],
    )
    replace_table_rows(
        document.tables[4],
        [
            ["序号", "字段名", "字段类型", "说明", "备注"],
            ["1", "schemaVersion", "INTEGER", "备份格式版本", "必填"],
            ["2", "exportedAt", "TEXT", "导出时间", "ISO 8601"],
            ["3", "tasks", "ARRAY", "任务数据集合", "非空"],
        ],
    )
    for table in document.tables:
        repeat_table_header(table)

    paragraphs = list(document.paragraphs)
    block_start = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip() == "4.5 功能模块划分"
    )
    block_anchor = next(
        paragraph
        for paragraph in paragraphs
        if paragraph.text.strip() == "图1 任务清单应用总体架构"
    )
    body = document._element.body
    children = list(body)
    start_index = children.index(block_start._element)
    end_index = children.index(document.tables[4]._element)
    for element in children[start_index : end_index + 1]:
        block_anchor._element.addprevious(element)

    references = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("[1] Android Developers")
    )
    set_paragraph_text(
        references,
        "\n".join(
            [
                "[1] Android Developers. Guide to app architecture.",
                "[2] Android Developers. Save data in a local database using Room.",
                "[3] Android Developers. Schedule tasks with WorkManager.",
                "[4] Kotlin Documentation. Coroutines guide.",
            ]
        ),
        first_line_indent=None,
    )


def polish_reference_list(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith("[1]") or "[2]" not in text:
            continue
        set_paragraph_text(
            paragraph,
            re.sub(r"\s+(?=\[\d+\])", "\n", text),
            first_line_indent=None,
        )


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--catalog", type=Path, required=True)
    parser.add_argument("--example-id", required=True)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    catalog = json.loads(args.catalog.read_text(encoding="utf-8"))
    example = next(
        item for item in catalog["examples"] if item["id"] == args.example_id
    )
    document = Document(args.input)
    if "dataTableRows" in example:
        fill_data_table(document, example["dataTableRows"])
    metadata_path = args.catalog.resolve().parents[2] / example["metadata"]
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    if example["id"] == "neutral-review-panel-lab":
        fill_review_score(document, str(metadata["成绩"]))
    if example["id"] == "neutral-code-notebook-lab":
        polish_code_notebook(document)
    if example["id"] == "neutral-project-dossier":
        polish_project_dossier(document)
    polish_reference_list(document)

    document.core_properties.author = str(metadata.get("姓名", ""))
    document.core_properties.last_modified_by = str(metadata.get("姓名", ""))
    document.core_properties.title = f"{example['displayName']} - {example['topic']}"
    document.core_properties.subject = str(metadata.get("课程名称", "课程实验报告"))
    document.core_properties.keywords = "实验报告, 课程设计, DOCX 模板"

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
