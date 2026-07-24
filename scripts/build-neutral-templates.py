from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


INK = "000000"
MUTED = "666666"
BLUE = "285B7A"
TEAL = "176B6D"
LIGHT_BLUE = "EAF1F5"
LIGHT_TEAL = "E8F3F1"
LIGHT_GRAY = "F2F4F6"
LIGHT_WARM = "F5F1EA"
LIGHT_GREEN = "EDF4EF"
CHARCOAL = "263238"
NAVY = "203A5F"
FOREST = "315B45"
WARM_INK = "4B4035"
WHITE = "FFFFFF"


def set_run_font(
    run,
    *,
    east_asia: str,
    ascii_font: str = "Times New Roman",
    size: float,
    bold: bool = False,
    color: str = "000000",
) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(
    style,
    *,
    east_asia: str,
    ascii_font: str,
    size: float,
    bold: bool = False,
    color: str = "000000",
) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color: str = "666666", size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def set_table_borders(table, *, color: str = "666666", size: int = 8) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_twips: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = sum(widths_twips)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_twips[min(index, len(widths_twips) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_page_border(section, *, color: str = "404040", size: int = 12, offset: int = 18) -> None:
    sect_pr = section._sectPr
    borders = sect_pr.find(qn("w:pgBorders"))
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        sect_pr.append(borders)
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), str(offset))
        element.set(qn("w:color"), color)


def add_page_number(paragraph, *, color: str = MUTED) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, east_asia="宋体", size=9, color=color)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=9, color=color)
    run._r.extend((begin, instruction, separate, value, end))
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, east_asia="宋体", size=9, color=color)


def configure_document(
    doc: Document,
    *,
    body_font: str,
    body_size: float,
    heading_font: str,
    heading_color: str,
    margins_cm: tuple[float, float, float, float] = (2.2, 2.2, 2.2, 2.4),
) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(margins_cm[0])
    section.right_margin = Cm(margins_cm[1])
    section.bottom_margin = Cm(margins_cm[2])
    section.left_margin = Cm(margins_cm[3])
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    set_style_font(
        normal,
        east_asia=body_font,
        ascii_font="Times New Roman",
        size=body_size,
        color=INK,
    )
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Pt(body_size * 2)

    for name, size, before, after in (
        ("Heading 1", 15, 12, 6),
        ("Heading 2", 12, 8, 4),
        ("Heading 3", 11, 6, 3),
    ):
        style = doc.styles[name]
        set_style_font(
            style,
            east_asia=heading_font,
            ascii_font="Arial",
            size=size,
            bold=True,
            color=heading_color,
        )
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Pt(0)

    caption = doc.styles["Caption"]
    set_style_font(
        caption,
        east_asia="宋体",
        ascii_font="Times New Roman",
        size=10.5,
        color=INK,
    )
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True

    properties = doc.core_properties
    properties.title = "中性实验报告模板"
    properties.subject = "通用实验报告与课程设计报告"
    properties.author = "experiment-report-skill"
    properties.last_modified_by = "experiment-report-skill"
    properties.keywords = "neutral, experiment report, course design"
    properties.comments = "Original neutral template generated by experiment-report-skill."


def add_title(
    doc: Document,
    text: str,
    *,
    font: str,
    size: float,
    color: str = "000000",
    subtitle: str | None = None,
    subtitle_color: str = MUTED,
    before: float = 2,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=font, ascii_font="Arial", size=size, bold=True, color=color)
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.first_line_indent = Pt(0)
        sub.paragraph_format.space_before = Pt(0)
        sub.paragraph_format.space_after = Pt(12)
        run = sub.add_run(subtitle)
        set_run_font(
            run,
            east_asia="宋体",
            ascii_font="Times New Roman",
            size=9.5,
            bold=False,
            color=subtitle_color,
        )


def add_metadata_table(
    doc: Document,
    *,
    rows: list[list[str]],
    widths_twips: list[int],
    label_font: str,
    value_font: str,
    label_fill: str | None = None,
    border_color: str = "666666",
    border_size: int = 8,
) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths_twips)
    set_table_borders(table, color=border_color, size=border_size)
    for row_index, values in enumerate(rows):
        for column_index, text in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell, color=border_color, size=border_size)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(text)
            is_label = column_index % 2 == 0
            if is_label:
                set_cell_no_wrap(cell)
            set_run_font(
                run,
                east_asia=label_font if is_label else value_font,
                ascii_font="Arial" if is_label else "Times New Roman",
                size=10.5 if is_label else 10.5,
                bold=is_label,
                color=INK,
            )
            if is_label and label_fill:
                set_cell_fill(cell, label_fill)
    after = doc.add_paragraph()
    after.paragraph_format.first_line_indent = Pt(0)
    after.paragraph_format.space_after = Pt(0)


def add_table_cell_text(
    cell,
    text: str,
    *,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size: float = 10.5,
    bold: bool = False,
    color: str = INK,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    fill: str | None = None,
    no_wrap: bool = False,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    if fill:
        set_cell_fill(cell, fill)
    if no_wrap:
        set_cell_no_wrap(cell)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    set_run_font(
        run,
        east_asia=east_asia,
        ascii_font=ascii_font,
        size=size,
        bold=bold,
        color=color,
    )


def add_sections(
    doc: Document,
    titles: list[str],
    *,
    heading_style: str = "Heading 1",
    body_font: str = "宋体",
    body_size: float = 12,
    placeholder_lines: int = 1,
) -> None:
    for title in titles:
        heading = doc.add_paragraph(style=heading_style)
        heading.paragraph_format.first_line_indent = Pt(0)
        heading.add_run(title)
        for _ in range(placeholder_lines):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(body_size * 2)
            paragraph.paragraph_format.keep_together = True
            run = paragraph.add_run()
            set_run_font(run, east_asia=body_font, size=body_size, color=INK)


def add_quiet_header_footer(doc: Document, label: str, *, color: str = MUTED) -> None:
    seen_headers: set[int] = set()
    seen_footers: set[int] = set()
    for section in doc.sections:
        header = section.header
        header_key = id(header._element)
        if header_key not in seen_headers:
            seen_headers.add(header_key)
            paragraph = header.paragraphs[0]
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            run = paragraph.add_run(label)
            set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=9, color=color)

        footer = section.footer
        footer_key = id(footer._element)
        if footer_key not in seen_footers:
            seen_footers.add(footer_key)
            paragraph = footer.paragraphs[0]
            paragraph.clear()
            add_page_number(paragraph, color=color)


def metadata_rows(experiment_label: str = "实验名称") -> list[list[str]]:
    return [
        ["姓名", "", "学号", ""],
        ["班级", "", "指导教师", ""],
        ["课程名称", "", experiment_label, ""],
        ["实验时间", "", "实验地点", ""],
        ["实验性质", "", "成绩", ""],
    ]


def build_classic(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=12,
        heading_font="黑体",
        heading_color="000000",
    )
    add_title(doc, "实验报告", font="黑体", size=22)
    add_metadata_table(
        doc,
        rows=metadata_rows(),
        widths_twips=[1500, 3300, 1500, 3300],
        label_font="黑体",
        value_font="宋体",
        label_fill=LIGHT_GRAY,
    )
    add_sections(
        doc,
        [
            "一、实验目的",
            "二、实验环境",
            "三、实验原理",
            "四、实验步骤",
            "五、实验结果",
            "六、问题分析",
            "七、实验总结",
        ],
    )
    add_quiet_header_footer(doc, "实验报告")
    doc.save(path)


def build_bordered(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="仿宋",
        body_size=12,
        heading_font="黑体",
        heading_color="000000",
        margins_cm=(2.35, 2.35, 2.35, 2.35),
    )
    set_page_border(doc.sections[0], color="303030", size=14, offset=20)
    add_title(doc, "实验报告", font="黑体", size=20)
    add_metadata_table(
        doc,
        rows=metadata_rows(),
        widths_twips=[1450, 3250, 1450, 3250],
        label_font="黑体",
        value_font="仿宋",
        border_color="303030",
        border_size=10,
    )
    add_sections(
        doc,
        [
            "一、实验任务",
            "二、实验条件",
            "三、实验原理",
            "四、实验过程",
            "五、实验记录",
            "六、问题处理",
            "七、结论与体会",
        ],
        body_font="仿宋",
        placeholder_lines=2,
    )
    add_quiet_header_footer(doc, "实验记录")
    doc.save(path)


def build_engineering(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=11.5,
        heading_font="黑体",
        heading_color=INK,
        margins_cm=(2.0, 2.0, 2.1, 2.2),
    )
    add_title(
        doc,
        "工程技术实验报告",
        font="黑体",
        size=21,
        color=INK,
        subtitle="TECHNICAL LAB REPORT",
    )
    add_metadata_table(
        doc,
        rows=metadata_rows(),
        widths_twips=[1500, 3300, 1500, 3300],
        label_font="黑体",
        value_font="宋体",
        label_fill=LIGHT_GRAY,
        border_color="707070",
    )
    add_sections(
        doc,
        [
            "1 实验目标与验收条件",
            "2 环境、工具与版本",
            "3 原理与技术方案",
            "4 实施步骤与关键配置",
            "5 测试数据与运行结果",
            "6 故障定位、结果分析与改进",
            "7 实验结论",
        ],
        body_size=11.5,
    )
    code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(
        code_style,
        east_asia="等线",
        ascii_font="Consolas",
        size=9.5,
        color=INK,
    )
    code_style.paragraph_format.first_line_indent = Pt(0)
    code_style.paragraph_format.left_indent = Cm(0.6)
    code_style.paragraph_format.right_indent = Cm(0.4)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    caption_style = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(
        caption_style,
        east_asia="宋体",
        ascii_font="Times New Roman",
        size=10,
        color=INK,
    )
    caption_style.paragraph_format.first_line_indent = Pt(0)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_quiet_header_footer(doc, "工程技术实验报告", color=MUTED)
    doc.save(path)


def build_course_design(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=12,
        heading_font="黑体",
        heading_color=INK,
        margins_cm=(2.4, 2.3, 2.4, 2.6),
    )
    cover = doc.sections[0]
    cover.different_first_page_header_footer = True
    add_title(
        doc,
        "课程设计报告",
        font="黑体",
        size=26,
        color=INK,
        subtitle="COURSE DESIGN REPORT",
        before=72,
    )
    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic.paragraph_format.first_line_indent = Pt(0)
    topic.paragraph_format.space_before = Pt(18)
    topic.paragraph_format.space_after = Pt(30)
    run = topic.add_run("课题名称：")
    set_run_font(run, east_asia="宋体", size=14, bold=True, color=INK)
    add_metadata_table(
        doc,
        rows=[
            ["学生姓名", ""],
            ["学号", ""],
            ["班级", ""],
            ["课程名称", ""],
            ["指导教师", ""],
            ["完成时间", ""],
        ],
        widths_twips=[2400, 6500],
        label_font="黑体",
        value_font="宋体",
        label_fill=LIGHT_GRAY,
        border_color="707070",
    )
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Mm(210)
    body_section.page_height = Mm(297)
    body_section.top_margin = Cm(2.3)
    body_section.right_margin = Cm(2.3)
    body_section.bottom_margin = Cm(2.4)
    body_section.left_margin = Cm(2.6)
    add_sections(
        doc,
        [
            "一、摘要",
            "二、关键词",
            "三、设计目标",
            "四、开发环境",
            "五、需求分析",
            "六、方案设计与实现",
            "七、实现结果",
            "八、问题与改进",
            "九、设计总结",
            "十、参考文献",
        ],
        placeholder_lines=1,
    )
    add_quiet_header_footer(doc, "课程设计报告")
    doc.save(path)


def build_modern(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="微软雅黑",
        body_size=11,
        heading_font="微软雅黑",
        heading_color=TEAL,
        margins_cm=(2.1, 2.1, 2.2, 2.3),
    )
    add_title(
        doc,
        "实验报告",
        font="微软雅黑",
        size=24,
        color=TEAL,
        subtitle="OBJECTIVE · PROCESS · EVIDENCE · CONCLUSION",
        subtitle_color=MUTED,
    )
    add_metadata_table(
        doc,
        rows=[
            ["实验名称", ""],
            ["课程名称", ""],
            ["姓名", ""],
            ["学号", ""],
            ["班级", ""],
            ["指导教师", ""],
        ],
        widths_twips=[2100, 6900],
        label_font="微软雅黑",
        value_font="微软雅黑",
        label_fill=LIGHT_TEAL,
        border_color="7EA9A5",
    )
    add_sections(
        doc,
        [
            "01 目标与范围",
            "02 环境与材料",
            "03 原理与方案",
            "04 方法与步骤",
            "05 结果与证据",
            "06 分析与修正",
            "07 总结",
        ],
        body_font="微软雅黑",
        body_size=11,
    )
    add_quiet_header_footer(doc, "实验报告", color=TEAL)
    doc.save(path)


def build_compact_header(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=11.5,
        heading_font="黑体",
        heading_color=INK,
        margins_cm=(1.8, 1.9, 2.0, 2.1),
    )
    add_title(
        doc,
        "实验记录与分析",
        font="黑体",
        size=20,
        color=INK,
        subtitle="COMPACT LAB RECORD",
        before=0,
    )
    rows = [
        ["课程名称", "", "实验名称", "", "日期", ""],
        ["姓名", "", "学号", "", "班级", ""],
        ["实验性质", "", "指导教师", "", "实验地点", ""],
    ]
    table = doc.add_table(rows=len(rows), cols=6)
    widths = [1150, 1800, 1150, 2250, 1050, 1800]
    set_table_geometry(table, widths)
    set_table_borders(table, color="555555", size=8)
    for row_index, values in enumerate(rows):
        for column_index, text in enumerate(values):
            cell = table.cell(row_index, column_index)
            is_label = column_index % 2 == 0
            add_table_cell_text(
                cell,
                text,
                east_asia="黑体" if is_label else "宋体",
                ascii_font="Arial" if is_label else "Times New Roman",
                size=9.5 if is_label else 10,
                bold=is_label,
                fill=LIGHT_GRAY if is_label else None,
                no_wrap=is_label,
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    add_sections(
        doc,
        [
            "一、实验目的",
            "二、实验环境",
            "三、实验原理",
            "四、实验步骤",
            "五、实验结果",
            "六、问题分析",
            "七、实验总结",
        ],
        body_size=11.5,
    )
    add_quiet_header_footer(doc, "实验记录与分析")
    doc.save(path)


def build_review_panel(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="仿宋",
        body_size=11.5,
        heading_font="黑体",
        heading_color=WARM_INK,
        margins_cm=(2.2, 2.2, 2.2, 2.35),
    )
    set_page_border(doc.sections[0], color="554A3F", size=12, offset=20)
    add_title(
        doc,
        "实验报告与评阅记录",
        font="黑体",
        size=20,
        color=WARM_INK,
    )
    table = doc.add_table(rows=4, cols=5)
    widths = [1400, 2650, 1400, 2650, 1100]
    set_table_geometry(table, widths)
    set_table_borders(table, color="6B6055", size=9)
    values = [
        ["课程名称", "", "实验性质", "", "成绩"],
        ["实验名称", "", "日期", "", ""],
        ["姓名", "", "学号", "", ""],
        ["班级", "", "指导教师", "", ""],
    ]
    score_cell = table.cell(0, 4).merge(table.cell(3, 4))
    for row in table.rows:
        set_row_cant_split(row)
    for row_index in range(4):
        for column_index in range(4):
            is_label = column_index % 2 == 0
            add_table_cell_text(
                table.cell(row_index, column_index),
                values[row_index][column_index],
                east_asia="黑体" if is_label else "仿宋",
                ascii_font="Arial" if is_label else "Times New Roman",
                size=10,
                bold=is_label,
                fill=LIGHT_WARM if is_label else None,
                no_wrap=is_label,
            )
    add_table_cell_text(
        score_cell,
        "成绩",
        east_asia="黑体",
        ascii_font="Arial",
        size=10,
        bold=True,
        fill=LIGHT_WARM,
        no_wrap=True,
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    add_sections(
        doc,
        [
            "一、实验目的",
            "二、实验条件",
            "三、实验原理",
            "四、实验步骤",
            "五、实验结果",
            "六、问题分析",
            "七、实验小结",
        ],
        body_font="仿宋",
        body_size=11.5,
        placeholder_lines=1,
    )
    review = doc.add_table(rows=2, cols=2)
    set_table_geometry(review, [1800, 7400])
    set_table_borders(review, color="6B6055", size=9)
    for row in review.rows:
        set_row_cant_split(row)
    add_table_cell_text(
        review.cell(0, 0),
        "教师评语",
        east_asia="黑体",
        ascii_font="Arial",
        size=10.5,
        bold=True,
        fill=LIGHT_WARM,
        no_wrap=True,
    )
    add_table_cell_text(review.cell(0, 1), "", east_asia="仿宋", size=10.5)
    add_table_cell_text(
        review.cell(1, 0),
        "签名与日期",
        east_asia="黑体",
        ascii_font="Arial",
        size=10.5,
        bold=True,
        fill=LIGHT_WARM,
        no_wrap=True,
    )
    add_table_cell_text(review.cell(1, 1), "", east_asia="仿宋", size=10.5)
    add_quiet_header_footer(doc, "实验报告与评阅记录", color="6B6055")
    doc.save(path)


def build_code_notebook(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=11.5,
        heading_font="黑体",
        heading_color=NAVY,
        margins_cm=(2.0, 2.0, 2.1, 2.2),
    )
    add_title(
        doc,
        "程序设计实验报告",
        font="黑体",
        size=22,
        color=NAVY,
        subtitle="IMPLEMENTATION · TEST · DEBUG",
        subtitle_color="52677E",
    )
    add_metadata_table(
        doc,
        rows=[
            ["课程名称", "", "实验名称", ""],
            ["姓名", "", "学号", ""],
            ["班级", "", "实验性质", ""],
            ["指导教师", "", "日期", ""],
        ],
        widths_twips=[1500, 3300, 1500, 3300],
        label_font="黑体",
        value_font="宋体",
        label_fill=LIGHT_BLUE,
        border_color="647A91",
    )
    heading_one = doc.styles["Heading 1"]
    heading_one.font.size = Pt(13)
    heading_one.paragraph_format.space_before = Pt(7)
    heading_one.paragraph_format.space_after = Pt(2)
    code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(
        code_style,
        east_asia="等线",
        ascii_font="Consolas",
        size=9.5,
        color=CHARCOAL,
    )
    code_style.paragraph_format.first_line_indent = Pt(0)
    code_style.paragraph_format.left_indent = Cm(0.45)
    code_style.paragraph_format.right_indent = Cm(0.35)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.keep_together = True
    command_style = doc.styles.add_style("Command Output", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(
        command_style,
        east_asia="等线",
        ascii_font="Consolas",
        size=9.5,
        color="23352A",
    )
    command_style.paragraph_format.first_line_indent = Pt(0)
    command_style.paragraph_format.left_indent = Cm(0.45)
    command_style.paragraph_format.right_indent = Cm(0.35)
    command_style.paragraph_format.space_before = Pt(4)
    command_style.paragraph_format.space_after = Pt(4)
    for index, title in enumerate(
        [
            "1 实验目的",
            "2 实验环境",
            "3 实验原理",
            "4 实验步骤",
            "5 实验结果",
            "6 问题分析",
            "7 实验总结",
        ]
    ):
        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.first_line_indent = Pt(0)
        heading.add_run(title)
        paragraph_style = "Code Block" if index == 3 else "Command Output" if index == 4 else None
        paragraph = doc.add_paragraph(style=paragraph_style)
        paragraph.paragraph_format.keep_together = True
        if paragraph_style:
            set_paragraph_shading(paragraph, "F3F5F7" if index == 3 else "F1F5F2")
    add_quiet_header_footer(doc, "程序设计实验报告", color=NAVY)
    doc.save(path)


def build_data_analysis(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=11.5,
        heading_font="黑体",
        heading_color=FOREST,
        margins_cm=(2.0, 2.1, 2.2, 2.25),
    )
    add_title(
        doc,
        "实验数据与分析报告",
        font="黑体",
        size=22,
        color=FOREST,
        subtitle="DATA · OBSERVATION · ANALYSIS",
        subtitle_color="5B7464",
    )
    add_metadata_table(
        doc,
        rows=[
            ["实验名称", "", "课程名称", ""],
            ["姓名", "", "学号", ""],
            ["班级", "", "日期", ""],
            ["实验性质", "", "指导教师", ""],
        ],
        widths_twips=[1550, 3250, 1550, 3250],
        label_font="黑体",
        value_font="宋体",
        label_fill=LIGHT_GREEN,
        border_color="6F8977",
    )
    heading_one = doc.styles["Heading 1"]
    heading_one.font.size = Pt(13)
    heading_one.paragraph_format.space_before = Pt(7)
    heading_one.paragraph_format.space_after = Pt(2)
    for index, title in enumerate(
        [
            "一、实验目的",
            "二、实验环境",
            "三、实验原理",
            "四、实验步骤",
            "五、实验结果",
            "六、问题分析",
            "七、实验总结",
        ]
    ):
        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.first_line_indent = Pt(0)
        heading.add_run(title)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(23)
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run()
        set_run_font(run, east_asia="宋体", size=11.5, color=INK)
        if index == 4:
            data_table = doc.add_table(rows=4, cols=5)
            widths = [850, 2400, 1550, 1550, 2850]
            set_table_geometry(data_table, widths)
            set_table_borders(data_table, color="78907E", size=8)
            headers = ["序号", "观测项", "原始值", "处理值", "备注"]
            set_repeat_table_header(data_table.rows[0])
            for row in data_table.rows:
                set_row_cant_split(row)
            for column_index, header in enumerate(headers):
                add_table_cell_text(
                    data_table.cell(0, column_index),
                    header,
                    east_asia="黑体",
                    ascii_font="Arial",
                    size=9.5,
                    bold=True,
                    fill=LIGHT_GREEN,
                    no_wrap=True,
                )
            for row_index in range(1, 4):
                for column_index in range(5):
                    add_table_cell_text(
                        data_table.cell(row_index, column_index),
                        "",
                        east_asia="宋体",
                        size=9.5,
                    )
            after = doc.add_paragraph()
            after.paragraph_format.first_line_indent = Pt(0)
            after.paragraph_format.space_after = Pt(0)
    add_quiet_header_footer(doc, "实验数据与分析报告", color=FOREST)
    doc.save(path)


def build_project_dossier(path: Path) -> None:
    doc = Document()
    configure_document(
        doc,
        body_font="宋体",
        body_size=12,
        heading_font="黑体",
        heading_color=CHARCOAL,
        margins_cm=(2.5, 2.4, 2.5, 2.7),
    )
    cover = doc.sections[0]
    cover.different_first_page_header_footer = True
    add_title(
        doc,
        "项目技术报告",
        font="黑体",
        size=27,
        color=CHARCOAL,
        subtitle="PROJECT TECHNICAL DOSSIER",
        subtitle_color="667078",
        before=82,
    )
    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic.paragraph_format.first_line_indent = Pt(0)
    topic.paragraph_format.space_before = Pt(22)
    topic.paragraph_format.space_after = Pt(30)
    run = topic.add_run("课题名称：")
    set_run_font(run, east_asia="宋体", size=15, bold=True, color=CHARCOAL)
    add_metadata_table(
        doc,
        rows=[
            ["学生姓名", ""],
            ["学号", ""],
            ["班级", ""],
            ["课程名称", ""],
            ["指导教师", ""],
            ["完成时间", ""],
        ],
        widths_twips=[2300, 6600],
        label_font="黑体",
        value_font="宋体",
        label_fill=LIGHT_GRAY,
        border_color="707980",
    )
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Mm(210)
    body.page_height = Mm(297)
    body.top_margin = Cm(2.3)
    body.right_margin = Cm(2.4)
    body.bottom_margin = Cm(2.5)
    body.left_margin = Cm(2.7)
    heading_one = doc.styles["Heading 1"]
    heading_one.font.size = Pt(13)
    heading_one.paragraph_format.space_before = Pt(7)
    heading_one.paragraph_format.space_after = Pt(2)
    add_title(
        doc,
        "项目技术报告",
        font="黑体",
        size=20,
        color=CHARCOAL,
        before=0,
    )
    add_sections(
        doc,
        [
            "摘要",
            "关键词",
            "1 设计目标",
            "2 开发环境",
            "3 需求分析",
            "4 方案设计与实现",
            "5 运行结果",
            "6 问题与改进",
            "7 设计总结",
            "参考文献",
        ],
        placeholder_lines=1,
    )
    add_quiet_header_footer(doc, "项目技术报告", color="5F6971")
    doc.save(path)


BUILDERS = {
    "neutral-classic-lab.docx": build_classic,
    "neutral-bordered-lab.docx": build_bordered,
    "neutral-engineering-lab.docx": build_engineering,
    "neutral-course-design.docx": build_course_design,
    "neutral-modern-minimal.docx": build_modern,
    "neutral-compact-header-lab.docx": build_compact_header,
    "neutral-review-panel-lab.docx": build_review_panel,
    "neutral-code-notebook-lab.docx": build_code_notebook,
    "neutral-data-analysis-lab.docx": build_data_analysis,
    "neutral-project-dossier.docx": build_project_dossier,
}


def build_all(output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    built: list[Path] = []
    for filename, builder in BUILDERS.items():
        path = output_dir / filename
        builder(path)
        built.append(path)
    shutil.copy2(output_dir / "neutral-classic-lab.docx", output_dir / "experiment-report-template.docx")
    shutil.copy2(output_dir / "neutral-course-design.docx", output_dir / "course-design-report-template.docx")
    return built


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir",
        default=str(Path(__file__).resolve().parents[1] / "examples" / "report-templates"),
    )
    args = parser.parse_args()
    for path in build_all(Path(args.output_dir).resolve()):
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
